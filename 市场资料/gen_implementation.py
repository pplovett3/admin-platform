# -*- coding: utf-8 -*-
"""
落地实施方案 - AI数字人工匠导师智能体
昌吉州AI智能体创新应用大赛
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ============ 样式设置 ============
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 页边距
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.8); s.right_margin = Cm(2.8)

C_PRIMARY = RGBColor(0x0E, 0x4D, 0xC4)
C_ORANGE = RGBColor(0xE8, 0x6A, 0x1A)
C_DARK = RGBColor(0x1A, 0x22, 0x33)
C_GRAY = RGBColor(0x5A, 0x6A, 0x7A)

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('a:shd')
    shd.set(qn('a:val'), 'clear')
    shd.set(qn('a:color'), 'auto')
    shd.set(qn('a:fill'), color_hex)
    tcPr.append(shd)

def add_heading(text, level=1, color=C_PRIMARY):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = '微软雅黑'
    run.font.color.rgb = color
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(13)
    return h

def add_para(text, size=11, bold=False, color=C_DARK, indent=True, spacing=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_bullet(text, size=11, color=C_DARK):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def style_table(table, header_color='0E4DC4'):
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头着色
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(10.5)
                r.font.name = '微软雅黑'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def fill_cell(cell, text, size=10.5, bold=False, color=C_DARK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ============ 封面 ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
run = p.add_run('落 地 实 施 方 案')
run.font.size = Pt(36); run.font.bold = True; run.font.color.rgb = C_PRIMARY
run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('AI 数字人工匠导师智能体')
run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = C_DARK
run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('职业教育全栈元宇宙解决方案')
run.font.size = Pt(16); run.font.color.rgb = C_GRAY
run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for _ in range(6):
    doc.add_paragraph()

info = [
    ('参赛赛事', '昌吉州 AI 智能体创新应用大赛'),
    ('参赛方向', 'AI + 职业教育数字化实训'),
    ('编制日期', '2026 年 7 月'),
    ('文档版本', 'V1.0'),
]
t = doc.add_table(rows=len(info), cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info):
    fill_cell(t.rows[i].cells[0], k, size=12, bold=True, color=C_PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(t.rows[i].cells[1], v, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    t.rows[i].cells[0].width = Cm(4)
    t.rows[i].cells[1].width = Cm(8)

doc.add_page_break()

# ============ 目录页 ============
add_heading('目  录', level=1)
toc_items = [
    '一、项目概述',
    '二、实施目标与原则',
    '三、总体实施路线',
    '四、实施时间表',
    '五、资源预算',
    '六、风险应对',
    '七、组织保障与分工',
    '八、质量管控与验收',
    '九、可持续发展机制',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(20)
    run = p.add_run(item)
    run.font.size = Pt(13); run.font.color.rgb = C_DARK
    run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ============ 一、项目概述 ============
add_heading('一、项目概述', level=1)

add_heading('1.1 项目背景', level=2)
add_para('本项目紧扣昌吉州 AI 智能体创新应用大赛主题，结合昌吉州作为国家"东数西算"枢纽、国际融合算力中心的定位，聚焦现代职业教育数字化实训的核心痛点立项。当前职业教育与企业技能培训中，3D 实训内容创作门槛高，非专业教师难以独立完成；虚拟仿真教学局限于单机体验，缺乏高并发低延迟的多人协同机制，且交互刻板无个性化指导。')

add_para('项目将 AI 技术与空间计算、WebGL2.0 渲染技术融合，打造 AI 数字人工匠导师智能体，既契合昌吉州算力资源优势，又能为本地职教和企业输送沉浸式实训平台，赋能数字经济与产业升级。')

add_heading('1.2 项目定位', level=2)
add_para('项目定位为"一中心 · 三角色 · 全栈闭环"的 AI 职教元宇宙解决方案：')
add_bullet('一中心：AI 工匠导师智能体，具备"感知-推理-执行"闭环能力；')
add_bullet('三角色：智能讲师（AI 数字人授课）、智能助教（个性化伴学）、智能教研员（零门槛 3D 课件创作）；')
add_bullet('全栈闭环：Web · Unity · VR 三端数据格式统一，创作-发布-学习-考核全流程闭环。')

add_heading('1.3 适用范围', level=2)
add_para('本方案适用于昌吉州及全疆职业院校数字化实训建设、企业技能内训、高危特种作业虚拟实操等场景，并可作为"一带一路"技能出海的数字化载体。')

# ============ 二、实施目标与原则 ============
add_heading('二、实施目标与原则', level=1)

add_heading('2.1 实施目标', level=2)
targets = [
    ('平台部署', '完成 AI 智能体平台在昌吉州智算中枢的部署，实现云-边-端协同架构落地'),
    ('场景落地', '遴选 2-3 所昌吉州职业院校作为试点，完成不少于 3 个专业的 3D 课件库建设'),
    ('能力验证', '验证 50 人高并发协同、100ms 内音频延迟、分钟级 3D 课件生成等核心指标'),
    ('算力消纳', '有效消纳昌吉州融合算力资源，形成"算力 + AI 职教"创新产业链雏形'),
    ('模式成型', '形成可复制、可推广的职教数字化模式，2027 年推广至 8 所院校，为全疆推广奠定基础'),
]
for i, (k, v) in enumerate(targets, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f'目标 {i}：{k} —— ')
    run.font.bold = True; run.font.color.rgb = C_ORANGE; run.font.size = Pt(11)
    run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run2 = p.add_run(v)
    run2.font.size = Pt(11); run2.font.color.rgb = C_DARK
    run2.font.name = '微软雅黑'; run2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_heading('2.2 实施原则', level=2)
principles = [
    ('算力优先', '优先依托昌吉州本地算力底座，推动算力与职教场景深度融合，体现"东数西算"枢纽价值'),
    ('试点先行', '以昌吉州本地职业院校为试点，验证模式后再规模化推广，降低实施风险'),
    ('安全合规', '建立算法伦理围栏与人机协同管控机制，保障数据安全与系统合规'),
    ('普惠开放', '全终端浏览器接入，无高端硬件依赖，确保偏远地区师生均可享受优质资源'),
    ('持续迭代', '采用敏捷开发与持续交付模式，根据教学反馈快速迭代产品能力'),
]
for k, v in principles:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f'● {k}：')
    run.font.bold = True; run.font.color.rgb = C_PRIMARY; run.font.size = Pt(11)
    run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run2 = p.add_run(v)
    run2.font.size = Pt(11); run2.font.color.rgb = C_DARK
    run2.font.name = '微软雅黑'; run2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ============ 三、总体实施路线 ============
add_heading('三、总体实施路线', level=1)
add_para('项目总体采用"三阶段递进"实施路线，从标杆试点到区域复制再到出海输出，逐步扩大应用规模与影响力。')

add_heading('3.1 第一阶段：标杆示范（2026 H2）', level=2)
add_para('以昌吉州本地职业院校为试点，完成平台部署与核心功能验证，打造区域数字职教标杆案例。重点完成云层智算中枢对接、3D 课件库建设、AI 智能体能力调优，并验证多人协同、低延迟通信等核心指标。')

add_heading('3.2 第二阶段：早期推广（2027）', level=2)
add_para('在标杆案例验证成功的基础上，推广至 8 所昌吉州及全疆职业院校，完善多专业场景库，形成标准化部署方案。同步建立"政校企"协同生态雏形，为后续规模化推广奠定基础。')

add_heading('3.3 第三阶段：规模复制与出海（2028+）', level=2)
add_para('在 8 所院校推广验证基础上，复制至全疆及西北更多职业院校，并依托新疆"一带一路"核心区定位，建设多语种实训平台，输出中国职业技能标准。')

# ============ 四、实施时间表 ============
add_heading('四、实施时间表', level=1)
add_para('以下为第一阶段（标杆示范阶段）的详细实施时间表，总周期 6 个月，分四个子阶段推进。')

add_heading('4.1 总体里程碑', level=2)
t = doc.add_table(rows=6, cols=4)
headers = ['阶段', '时间区间', '核心任务', '交付物']
for i, h in enumerate(headers):
    fill_cell(t.rows[0].cells[i], h, align=WD_ALIGN_PARAGRAPH.CENTER)
rows = [
    ('启动筹备', '第 1 月', '需求调研、试点院校遴选、团队组建、算力资源对接', '需求规格说明书、试点合作协议'),
    ('平台部署', '第 2-3 月', '云层智算中枢部署、微服务搭建、CDN 边缘节点配置、端侧渲染引擎集成', '平台部署完成报告、架构图'),
    ('内容建设', '第 4-5 月', '3D 课件库建设、AI 智能体调优、教师培训、课件审核发布', '课件库（≥3 专业）、培训记录'),
    ('试点验证', '第 6 月', '教学试运行、性能指标测试、用户反馈收集、案例总结', '测试报告、标杆案例报告'),
]
for i, row in enumerate(rows, 1):
    for j, val in enumerate(row):
        fill_cell(t.rows[i].cells[j], val, size=10, align=WD_ALIGN_PARAGRAPH.LEFT if j > 1 else WD_ALIGN_PARAGRAPH.CENTER)
style_table(t)

add_heading('4.2 详细任务时间表（甘特图数据）', level=2)
tasks = [
    ('T1', '需求调研与院校遴选', '第1月第1-2周', '产品经理', '完成'),
    ('T2', '算力资源对接与协议签署', '第1月第3-4周', '项目经理', '完成'),
    ('T3', '云层智算中枢环境部署', '第2月第1-2周', '运维团队', '关键里程碑'),
    ('T4', '微服务后端搭建与联调', '第2月第3-4周', '后端团队', '关键里程碑'),
    ('T5', 'CDN 边缘节点配置', '第3月第1周', '运维团队', '完成'),
    ('T6', 'Web 端渲染引擎集成', '第3月第2-3周', '前端团队', '关键里程碑'),
    ('T7', 'Unity/VR 客户端对接', '第3月第4周', '客户端团队', '完成'),
    ('T8', 'AI 大模型接入与调优', '第4月第1-2周', 'AI 团队', '关键里程碑'),
    ('T9', '3D 课件库建设（3 专业）', '第4月第3周-第5月第2周', '内容团队', '关键里程碑'),
    ('T10', '教师培训与课件审核', '第5月第3-4周', '教研团队', '完成'),
    ('T11', '教学试运行', '第6月第1-2周', '试点院校', '关键里程碑'),
    ('T12', '性能指标测试', '第6月第3周', '测试团队', '完成'),
    ('T13', '案例总结与推广准备', '第6月第4周', '项目经理', '完成'),
]
t = doc.add_table(rows=len(tasks)+1, cols=5)
headers = ['编号', '任务名称', '时间区间', '责任人', '状态']
for i, h in enumerate(headers):
    fill_cell(t.rows[0].cells[i], h, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row in enumerate(tasks, 1):
    for j, val in enumerate(row):
        fill_cell(t.rows[i].cells[j], val, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER if j in (0,4) else WD_ALIGN_PARAGRAPH.LEFT)
style_table(t)

doc.add_page_break()

# ============ 五、资源预算 ============
add_heading('五、资源预算', level=1)
add_para('本预算覆盖标杆示范阶段（2026 H2，6 个月）所需的人力、算力、内容及运营资源，标杆阶段总预算 90 万元。后续早期推广阶段（2027 年）追加 110 万元，两阶段合计 200 万元。人力成本按参赛核心团队兼职+部分全职综合测算。')

add_heading('5.1 人力资源预算', level=2)
hr = [
    ('项目经理', '1', '6', '1.0', '6.0', '统筹协调、进度管控（兼职）'),
    ('AI 算法工程师', '1', '6', '1.5', '9.0', '大模型接入、智能体调优'),
    ('后端工程师', '1', '6', '1.3', '7.8', '微服务、数据湖、API'),
    ('前端工程师', '2', '6', '1.2', '14.4', 'Web 渲染引擎、编辑器'),
    ('3D 内容工程师', '1', '4', '1.0', '4.0', '3D 课件制作、模型转码'),
    ('Unity/VR 工程师', '1', '4', '1.2', '4.8', 'Unity 客户端、VR 适配'),
    ('运维工程师', '1', '6', '0.8', '4.8', '部署、监控、扩容（兼职）'),
    ('教研顾问', '1', '3', '0.8', '2.4', '教学设计、课件审核（兼职）'),
    ('测试工程师', '1', '2', '1.0', '2.0', '性能测试、功能测试'),
]
t = doc.add_table(rows=len(hr)+2, cols=6)
headers = ['角色', '人数', '月数', '月薪(万)', '小计(万)', '职责']
for i, h in enumerate(headers):
    fill_cell(t.rows[0].cells[i], h, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row in enumerate(hr, 1):
    for j, val in enumerate(row):
        fill_cell(t.rows[i].cells[j], val, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER if j in (1,2,3,4) else WD_ALIGN_PARAGRAPH.LEFT)
# 合计行
fill_cell(t.rows[-1].cells[0], '合计', size=10.5, bold=True, color=C_ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
for j in range(1, 4):
    fill_cell(t.rows[-1].cells[j], '', size=10.5)
fill_cell(t.rows[-1].cells[4], '55.2', size=10.5, bold=True, color=C_ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
fill_cell(t.rows[-1].cells[5], '', size=10.5)
style_table(t)

add_heading('5.2 算力与基础设施预算', level=2)
infra = [
    ('昌吉州智算中枢算力租赁', 'GPU 推理实例 × 1', '6 月', '5.0', '30.0', 'DeepSeek/QwenVL 推理（含补贴优惠价）'),
    ('CDN 边缘节点', '带宽 50Mbps', '6 月', '0.6', '3.6', '3D 资产分发加速'),
    ('对象存储', '3 TB', '6 月', '0.15', '0.9', '3D 模型/课件存储'),
    ('域名与 SSL', '1 套', '1 年', '0.5', '0.5', 'HTTPS 反向代理'),
]
t = doc.add_table(rows=len(infra)+2, cols=6)
headers = ['项目', '规格', '周期', '单价(万)', '小计(万)', '说明']
for i, h in enumerate(headers):
    fill_cell(t.rows[0].cells[i], h, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row in enumerate(infra, 1):
    for j, val in enumerate(row):
        fill_cell(t.rows[i].cells[j], val, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER if j in (2,3,4) else WD_ALIGN_PARAGRAPH.LEFT)
fill_cell(t.rows[-1].cells[0], '合计', size=10.5, bold=True, color=C_ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
for j in range(1, 4):
    fill_cell(t.rows[-1].cells[j], '', size=10.5)
fill_cell(t.rows[-1].cells[4], '35.0', size=10.5, bold=True, color=C_ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
fill_cell(t.rows[-1].cells[5], '', size=10.5)
style_table(t)

add_heading('5.3 内容与运营费用预算', level=2)
other = [
    ('3D 模型与教研素材采购', '—', '—', '5.0', '5.0', '3D 模型、教学素材、教研资料'),
    ('试点院校合作经费', '3 所院校', '—', '1.0', '3.0', '试点部署配合、教师培训'),
    ('差旅与调研', '—', '6 月', '0.3', '1.8', '院校走访、需求调研'),
    ('市场推广与案例宣传', '—', '2 月', '0.6', '1.2', '标杆案例包装、宣传物料'),
]
t = doc.add_table(rows=len(other)+2, cols=6)
headers = ['项目', '数量', '周期', '单价(万)', '小计(万)', '说明']
for i, h in enumerate(headers):
    fill_cell(t.rows[0].cells[i], h, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row in enumerate(other, 1):
    for j, val in enumerate(row):
        fill_cell(t.rows[i].cells[j], val, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER if j in (1,2,3,4) else WD_ALIGN_PARAGRAPH.LEFT)
fill_cell(t.rows[-1].cells[0], '合计', size=10.5, bold=True, color=C_ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
for j in range(1, 4):
    fill_cell(t.rows[-1].cells[j], '', size=10.5)
fill_cell(t.rows[-1].cells[4], '11.0', size=10.5, bold=True, color=C_ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
fill_cell(t.rows[-1].cells[5], '', size=10.5)
style_table(t)

add_heading('5.4 预算汇总', level=2)
summary = [
    ('一、人力资源', '55.2', '标杆阶段 6 月'),
    ('二、算力与基础设施', '35.0', '标杆阶段 6 月'),
    ('三、内容与运营', '11.0', '标杆阶段 6 月'),
    ('标杆示范阶段小计', '101.2', '2026 H2（注1）'),
    ('早期推广阶段', '98.8', '2027 年'),
    ('总计', '200.0', '1.5 年'),
]
t = doc.add_table(rows=len(summary)+1, cols=3)
headers = ['预算类别', '金额(万元)', '说明']
for i, h in enumerate(headers):
    fill_cell(t.rows[0].cells[i], h, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row in enumerate(summary, 1):
    bold = (i >= 4)
    color = C_ORANGE if i == len(summary) else (C_PRIMARY if bold else C_DARK)
    for j, val in enumerate(row):
        fill_cell(t.rows[i].cells[j], val, size=11 if bold else 10, bold=bold, color=color, align=WD_ALIGN_PARAGRAPH.CENTER if j == 1 else WD_ALIGN_PARAGRAPH.LEFT)
style_table(t)
add_para('注1：标杆阶段三项合计约 101.2 万元，因部分岗位兼职折算及补贴优惠，实际支出以 90 万元为控制目标，差额通过参赛奖金及资源赞助弥补。早期推广阶段 98.8 万元含人力 50 万、算力 35 万、内容与运营 13.8 万。', spacing=6, color=C_GRAY)

doc.add_page_break()

# ============ 六、风险应对 ============
add_heading('六、风险应对', level=1)
add_para('项目实施过程中可能面临技术、市场、运营、合规等多维度风险，本节识别主要风险并制定针对性应对措施。')

add_heading('6.1 风险识别与应对矩阵', level=2)
risks = [
    ('R1', '技术风险', '高并发场景下渲染性能不达标，50 人同频出现卡顿',
     '高', '采用动态 LOD + 遮挡剔除 + 实例化渲染优化；预留压力测试阶段；配置弹性扩容机制', '中'),
    ('R2', '技术风险', 'AI 大模型推理延迟过高，影响交互体验',
     '中', '采用昌吉州本地算力就近推理；模型量化压缩；关键结果缓存；流式输出', '低'),
    ('R3', '技术风险', '弱网环境下 3D 资产加载失败',
     '中', 'CDN 边缘缓存 + 预加载策略；弱网降级方案；最低 2Mbps 可用保障', '低'),
    ('R4', '市场风险', '试点院校教师接受度低，3D 课件创作意愿不足',
     '中', '提供"文本→3D 课件"零门槛工具；安排专项培训；建立激励机制', '低'),
    ('R5', '市场风险', '昌吉州算力资源供给不足或价格波动',
     '中', '提前签署算力合作协议；预留备用云厂商；按需扩容控制成本', '低'),
    ('R6', '运营风险', '项目进度延期，关键里程碑无法按期达成',
     '中', '敏捷开发 + 周度进度跟踪；关键路径任务预留缓冲；建立升级预警机制', '低'),
    ('R7', '运营风险', '3D 课件内容质量不达标，不符合教学需求',
     '中', '教研顾问全程参与；建立课件审核流程；试点教师反馈迭代', '低'),
    ('R8', '合规风险', '数据安全与隐私合规问题',
     '高', '数据湖全生命周期治理；算法伦理围栏；人机协同管控；数据本地化存储', '中'),
    ('R9', '合规风险', 'AI 生成内容存在错误或不当信息',
     '中', '教师审核发布机制；AI 输出内容标注；建立纠错反馈通道', '低'),
    ('R10', '团队风险', '核心技术人员流失',
     '低', '关键技术文档化；岗位 AB 角；项目激励绑定；知识传承机制', '低'),
]
t = doc.add_table(rows=len(risks)+1, cols=6)
headers = ['编号', '类别', '风险描述', '风险等级', '应对措施', '残余风险']
for i, h in enumerate(headers):
    fill_cell(t.rows[0].cells[i], h, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row in enumerate(risks, 1):
    for j, val in enumerate(row):
        align = WD_ALIGN_PARAGRAPH.CENTER if j in (0,1,3,5) else WD_ALIGN_PARAGRAPH.LEFT
        color = C_ORANGE if row[3] == '高' else C_DARK
        fill_cell(t.rows[i].cells[j], val, size=9, color=color if j == 3 else C_DARK, align=align)
style_table(t)

add_heading('6.2 风险监控机制', level=2)
add_bullet('周度风险评审：每周项目例会同步风险状态，更新风险登记册；')
add_bullet('指标预警：设置性能、进度、成本三类预警阈值，触发后自动升级；')
add_bullet('应急响应：高风险事件 24 小时内启动应急响应，制定专项处置方案；')
add_bullet('复盘改进：每月进行风险复盘，将应对经验沉淀为标准流程。')

doc.add_page_break()

# ============ 七、组织保障与分工 ============
add_heading('七、组织保障与分工', level=1)

add_heading('7.1 组织架构', level=2)
add_para('项目设立三级组织架构：决策层、管理层、执行层，确保决策高效、执行有力。')

org = [
    ('决策层', '项目指导委员会', '由参赛团队负责人、昌吉州算力中心代表、试点院校领导组成，负责重大决策与资源协调'),
    ('管理层', '项目管理办公室（PMO）', '设项目经理 1 名，负责进度管控、风险管理、跨团队协调'),
    ('执行层', 'AI 算法组 / 后端组 / 前端组 / 内容组 / 运维组 / 测试组', '各组分组长负责本组任务执行与质量交付'),
]
t = doc.add_table(rows=len(org)+1, cols=3)
headers = ['层级', '组织', '职责']
for i, h in enumerate(headers):
    fill_cell(t.rows[0].cells[i], h, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row in enumerate(org, 1):
    for j, val in enumerate(row):
        fill_cell(t.rows[i].cells[j], val, size=10, align=WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT)
style_table(t)

add_heading('7.2 关键角色职责', level=2)
roles = [
    ('项目经理', '总体统筹、进度管控、风险协调、对外沟通、里程碑验收'),
    ('AI 算法组长', '大模型选型与接入、智能体能力调优、多模态交互实现'),
    ('后端组长', '微服务架构、数据湖建设、API 设计、弹性扩容'),
    ('前端组长', 'WebGL2.0 渲染引擎、3D 编辑器、播放器、WebXR 适配'),
    ('内容组长', '3D 课件制作、模型转码、教师培训、课件审核'),
    ('运维组长', '环境部署、CDN 配置、监控告警、安全保障'),
]
t = doc.add_table(rows=len(roles)+1, cols=2)
headers = ['角色', '核心职责']
for i, h in enumerate(headers):
    fill_cell(t.rows[0].cells[i], h, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row in enumerate(roles, 1):
    for j, val in enumerate(row):
        fill_cell(t.rows[i].cells[j], val, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT)
style_table(t)

# ============ 八、质量管控与验收 ============
add_heading('八、质量管控与验收', level=1)

add_heading('8.1 质量管控体系', level=2)
add_para('项目建立"三级质量管控"体系，确保交付物质量满足要求：')
add_bullet('一级：组内自测 —— 开发人员完成开发后进行自测，确保基本功能可用；')
add_bullet('二级：组间联调 —— 跨组联调测试，确保模块间接口与数据流通；')
add_bullet('三级：用户验收 —— 试点院校教师与学生在真实教学场景中验收。')

add_heading('8.2 验收标准', level=2)
criteria = [
    ('功能完整性', '六大核心模块全部上线，功能符合需求规格说明书'),
    ('性能指标', '50 人并发不卡顿；音频延迟<100ms；白板同步<50ms；Web 端 60FPS；VR 端 90FPS'),
    ('内容质量', '课件库覆盖≥3 个专业，每个专业≥10 个课件，教师审核通过率≥90%'),
    ('用户体验', '试点师生满意度评分≥4.0（满分 5 分）；关键操作响应≤200ms'),
    ('安全合规', '通过数据安全审计；算法伦理围栏生效；无高危漏洞'),
    ('文档完整', '部署手册、API 文档、用户手册、运维手册齐全'),
]
t = doc.add_table(rows=len(criteria)+1, cols=2)
headers = ['验收维度', '验收标准']
for i, h in enumerate(headers):
    fill_cell(t.rows[0].cells[i], h, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, row in enumerate(criteria, 1):
    for j, val in enumerate(row):
        fill_cell(t.rows[i].cells[j], val, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT)
style_table(t)

# ============ 九、可持续发展机制 ============
add_heading('九、可持续发展机制', level=1)

add_heading('9.1 运营可持续', level=2)
add_bullet('建立"平台 + 内容 + 服务"三位一体运营模式，平台订阅 + 内容分成 + 增值服务多元营收；')
add_bullet('通过规模化推广降低边际成本，实现轻量化零边际成本生产；')
add_bullet('建立教师创作者生态，激励教师贡献优质 3D 课件，形成内容飞轮。')

add_heading('9.2 技术可持续', level=2)
add_bullet('跟踪 AI 大模型、WebXR、空间计算前沿技术，持续迭代产品能力；')
add_bullet('建立模块化架构，新功能即插即用，避免技术债务积累；')
add_bullet('沉淀技术资产，形成可复用的组件库与工具链。')

add_heading('9.3 生态可持续', level=2)
add_bullet('构建"政校企"协同生态，政府政策引导、学校场景验证、企业技术赋能；')
add_bullet('联合昌吉州算力中心，打造"算力 + AI 职教"创新产业链；')
add_bullet('开放平台能力，吸引上下游合作伙伴共建职教元宇宙生态。')

# 结尾
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('— 本方案编制完毕 —')
run.font.size = Pt(12); run.font.color.rgb = C_GRAY
run.font.name = '微软雅黑'; run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

out = r"d:\Admin_Platform_Project\admin-platform\市场资料\落地实施方案-v2.docx"
doc.save(out)
print("落地实施方案已生成:", out)
